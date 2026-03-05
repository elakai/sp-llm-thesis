import os
import io
import json
import re
import cv2
import numpy as np
import fitz 
import pdfplumber
import pandas as pd
import docx
from PIL import Image
import pytesseract
from datetime import datetime
from typing import List, Dict, Tuple
from pinecone import Pinecone
from collections import defaultdict
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from src.config.settings import PINECONE_INDEX_NAME, get_vectorstore, TESSERACT_CMD
from src.config.constants import DOCS_FOLDER, MANIFEST_FILE, CHUNK_SIZE, CHUNK_OVERLAP
from src.config.logging_config import logger
from src.core.retrieval import invalidate_cache
from src.config.constants import VALID_CATEGORIES
from src.core.feedback import supabase


pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
logger.info(f"Tesseract CMD set to: {TESSERACT_CMD}")

def normalize_source_key(filename: str) -> str:
    return filename.strip().replace("\\", "/")

def is_already_ingested(filename: str) -> bool:
    """Checks Supabase manifest instead of local JSON."""
    norm_name = normalize_source_key(filename)
    try:
        response = supabase.table("manifest") \
            .select("status") \
            .eq("filename", norm_name) \
            .execute()
        return bool(response.data) and response.data[0]["status"] == "Active"
    except Exception as e:
        logger.error(f"Manifest check failed: {e}")
        return False

def clean_text(text: str) -> str:
    if not text: return ""
    text = re.sub(r'\n\s*\n', '\n\n', text) 
    return text.strip()

# --- VISION & TABLE PROCESSING ---
def preprocess_image_for_ocr(pil_image: Image) -> Image:
    """Gentle preprocessing: grayscale + Otsu binarization.
    Avoids aggressive adaptive thresholding that destroys word spacing.
    """
    img = np.array(pil_image.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(binary)

def post_process_ocr_text(text: str) -> str:
    """Fix common OCR artifacts: missing spaces between concatenated words."""
    if not text:
        return text
    # Space before uppercase after lowercase: "ofScience" -> "of Science"
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    # Space between letter and opening paren: "program(" -> "program ("
    text = re.sub(r'([a-zA-Z0-9])\(', r'\1 (', text)
    # Space between closing paren and letter: ")The" -> ") The"
    text = re.sub(r'\)([a-zA-Z])', r') \1', text)
    # Space after period/comma before uppercase: ".The" -> ". The"
    text = re.sub(r'([.,;:])([A-Z])', r'\1 \2', text)
    # Collapse multiple spaces
    text = re.sub(r' {2,}', ' ', text)
    return text

# Regex for curriculum section headers like "FIRST YEAR - First Semester", "Year 2, Semester 1"
_SECTION_HDR_RE = re.compile(
    r'(?:(?:FIRST|SECOND|THIRD|FOURTH|FIFTH)\s*YEAR|Year\s*\d+)'
    r'\s*[-\u2013\u2014,.:;\s]+\s*'
    r'(?:(?:1st|2nd|First|Second)\s*Semester|Semester\s*\d+|Summer|Intersession)',
    re.IGNORECASE,
)

def extract_program_info(filename: str) -> dict:
    """Parse program name and code from curriculum filename.
    e.g. 'CURRICULUM FOR BACHELOR OF SCIENCE IN ARCHITECTURE (BS ARCH).pdf'
    → {'program_full': 'BACHELOR OF SCIENCE IN ARCHITECTURE', 'program_code': 'BS ARCH'}
    """
    match = re.search(r'CURRICULUM\s+FOR\s+(.+?)\s*\(([^)]+)\)', filename, re.IGNORECASE)
    if match:
        return {"program_full": match.group(1).strip(), "program_code": match.group(2).strip()}
    return {}

def find_section_headers_for_tables(all_words: list, table_bboxes: list) -> dict:
    if not all_words or not table_bboxes:
        return {}
    # Group words into lines by y-coordinate
    lines: Dict[int, List[dict]] = {}
    for w in all_words:
        y = round(w['top'])
        merged = False
        for ey in list(lines.keys()):
            if abs(y - ey) < 5:
                lines[ey].append(w)
                merged = True
                break
        if not merged:
            lines[y] = [w]
    # Build sorted (y, text) list
    line_list = []
    for y in sorted(lines.keys()):
        words_in_line = sorted(lines[y], key=lambda w: w['x0'])
        text = ' '.join(w['text'] for w in words_in_line)
        line_list.append((y, text))
    # For each table, find the closest matching header above it
    result = {}
    for idx, bbox in enumerate(table_bboxes):
        table_top = bbox[1]
        best = None
        for y, text in line_list:
            if y >= table_top - 2:
                break
            if _SECTION_HDR_RE.search(text):
                best = text.strip()
        if best:
            result[idx] = best
    return result

def convert_table_to_markdown(table_data: list) -> str:
    if not table_data:
        return ""

    # Step 1: Forward-fill None values across each row (horizontal merges)
    filled = []
    for row in table_data:
        filled_row = []
        last_val = ""
        for cell in row:
            if cell is None or str(cell).strip() == "":
                filled_row.append(last_val)
            else:
                last_val = str(cell).replace('\n', ' ').strip()
                filled_row.append(last_val)
        filled.append(filled_row)

    if not filled:
        return ""

    # Step 2: Normalize column count
    max_cols = max(len(row) for row in filled)
    for row in filled:
        while len(row) < max_cols:
            row.append("")

    # Step 3: Down-fill None values down columns (vertical merges)
    for col_idx in range(max_cols):
        last_val = ""
        for row in filled:
            if row[col_idx] == "":
                row[col_idx] = last_val
            else:
                last_val = row[col_idx]

    # Step 4: Detect header rows (rows with no numeric data)
    def _is_header_row(row: list) -> bool:
        for cell in row:
            stripped = str(cell).replace('-', '').replace(' ', '').replace('–', '')
            try:
                float(stripped)
                return False
            except ValueError:
                continue
        return True

    header_rows = []
    data_start = 0
    for i, row in enumerate(filled):
        if _is_header_row(row):
            header_rows.append(row)
            data_start = i + 1
        else:
            break

    # Step 5: Merge multi-row headers into single descriptive column names
    if len(header_rows) > 1:
        merged_header = []
        for col_idx in range(max_cols):
            col_values = []
            for hrow in header_rows:
                val = hrow[col_idx] if col_idx < len(hrow) else ""
                if val and val not in col_values:
                    col_values.append(val)
            merged_header.append(" — ".join(col_values))
        header = merged_header
    elif header_rows:
        header = header_rows[0]
        data_start = 1
    else:
        header = filled[0]
        data_start = 1

    # Step 6: Build Markdown
    md = "| " + " | ".join(header) + " |\n"
    md += "| " + " | ".join(["---"] * len(header)) + " |\n"
    for row in filled[data_start:]:
        md += "| " + " | ".join(row) + " |\n"

    return md

def is_inside_any_bbox(word: dict, bboxes: list) -> bool:
    wx = (word['x0'] + word['x1']) / 2
    wy = (word['top'] + word['bottom']) / 2
    return any(x0 <= wx <= x1 and top <= wy <= bottom for (x0, top, x1, bottom) in bboxes)

def reconstruct_body_text(words: list) -> str:
    lines: Dict[int, List[dict]] = {}
    for word in words:
        found = False
        for y_coord in lines.keys():
            if abs(word['top'] - y_coord) < 5:
                lines[y_coord].append(word)
                found = True
                break
        if not found: lines[word['top']] = [word]
    sorted_y = sorted(lines.keys())
    text = ""
    for y in sorted_y:
        line_words = sorted(lines[y], key=lambda w: w['x0'])
        line_str = ""
        last_x1 = 0
        for w in line_words:
            if last_x1 > 0 and (w['x0'] - last_x1) > 5: line_str += " "
            line_str += w['text']
            last_x1 = w['x1']
        text += line_str + "\n"
    return text

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract ALL text from a DOCX file, including content in text boxes,
    shapes, and SmartArt that python-docx's doc.paragraphs misses.
    Parses the DOCX zip at the XML level for comprehensive extraction.
    """
    import zipfile
    from lxml import etree

    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    A_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'

    parts = []
    seen = set()

    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
        for entry in zf.namelist():
            if not entry.endswith('.xml'):
                continue
            # Main body, headers, footers, and SmartArt/diagram data
            if not any(entry.startswith(p) for p in [
                'word/document', 'word/header', 'word/footer', 'word/diagrams/data'
            ]):
                continue

            try:
                tree = etree.fromstring(zf.read(entry))
            except Exception:
                continue

            # WordprocessingML paragraphs (body text, tables, text boxes)
            for p in tree.iter(f'{W_NS}p'):
                runs = [t.text for t in p.iter(f'{W_NS}t') if t.text]
                line = ''.join(runs).strip()
                if line and line not in seen:
                    seen.add(line)
                    parts.append(line)

            # DrawingML text (SmartArt labels, shape text, charts)
            for t in tree.iter(f'{A_NS}t'):
                if t.text and t.text.strip() and t.text.strip() not in seen:
                    seen.add(t.text.strip())
                    parts.append(t.text.strip())

    return '\n'.join(parts)

# --- LOADERS ---
def load_pdf(path: str, filename: str) -> List[Document]:
    logger.info(f"Reading PDF: {filename}...")
    norm_filename = normalize_source_key(filename)
    program_info = extract_program_info(filename)
    docs = []

    try:
        fitz_doc = fitz.open(path)
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                fitz_page = fitz_doc[page_num - 1]

                # ── TABLE EXTRACTION ──────────────────────────────────────
                tables = page.find_tables()
                table_bboxes = [t.bbox for t in tables]

                for table in tables:
                    data = table.extract()
                    if not data:
                        continue
                    md = convert_table_to_markdown(data)
                    if not md.strip():
                        continue
                    prefix = (
                        f"Program: {program_info['program_full']} "
                        f"({program_info['program_code']})\n\n"
                        if program_info else ""
                    )
                    meta_t = {"source": norm_filename, "page": page_num, "type": "table"}
                    if program_info:
                        meta_t["program_code"] = program_info.get("program_code", "")
                    docs.append(Document(page_content=prefix + md, metadata=meta_t))

                # ── BODY TEXT (excluding table regions) ───────────────────
                all_words = page.extract_words(x_tolerance=3, y_tolerance=3)

                if table_bboxes:
                    non_table_words = [
                        w for w in all_words
                        if not is_inside_any_bbox(w, table_bboxes)
                    ]
                    body_text = clean_text(reconstruct_body_text(non_table_words))
                else:
                    # No tables — try fitz first, fall back to pdfplumber words
                    body_text = clean_text(fitz_page.get_text("text"))
                    if len(body_text) < 100:
                        plumber_text = clean_text(reconstruct_body_text(all_words))
                        if len(plumber_text) > len(body_text):
                            body_text = plumber_text

                # ── OCR FALLBACK (genuinely scanned pages only) ────────────
                if len(body_text) < 50 and not table_bboxes:
                    raw_fitz = fitz_page.get_text("text").strip()
                    if not raw_fitz or len(raw_fitz) < 50:
                        try:
                            logger.warning(f"Page {page_num} appears scanned. Running Tesseract OCR...")
                            pix = fitz_page.get_pixmap(dpi=300)
                            from PIL import Image as _PIL
                            import io as _io
                            pil_img = _PIL.open(_io.BytesIO(pix.tobytes("png")))
                            clean_img = preprocess_image_for_ocr(pil_img)
                            ocr_text = pytesseract.image_to_string(
                                clean_img, config="--psm 3 --oem 3"
                            )
                            body_text = clean_text(post_process_ocr_text(ocr_text))
                        except Exception as ocr_err:
                            logger.warning(f"OCR fallback failed on page {page_num}: {ocr_err}")

                if len(body_text) > 20:
                    prefix = (
                        f"Program: {program_info['program_full']} "
                        f"({program_info['program_code']})\n\n"
                        if program_info else ""
                    )
                    meta = {"source": norm_filename, "page": page_num, "type": "text"}
                    if program_info:
                        meta["program_code"] = program_info.get("program_code", "")
                    docs.append(Document(page_content=prefix + body_text, metadata=meta))

        fitz_doc.close()

    except Exception as e:
        logger.error(f"PDF Processing Error: {e}")

    return docs

def load_spreadsheet(path: str, filename: str, is_csv: bool = False) -> List[Document]:
    """
    Reads an Excel (.xlsx/.xls) or CSV file into a single Markdown table Document.
    Missing cells are filled with 'N/A'.  Uses ``tabulate`` via
    ``df.to_markdown()`` when available; falls back to ``convert_table_to_markdown``.
    Returns a list with one Document of type 'table'.
    """
    docs = []
    norm_filename = normalize_source_key(filename)
    try:
        df = pd.read_csv(path) if is_csv else pd.read_excel(path)
        df.fillna("N/A", inplace=True)
        try:
            md_table = df.to_markdown(index=False)
        except ImportError:
            md_table = convert_table_to_markdown([df.columns.tolist()] + df.values.tolist())
        docs.append(Document(
            page_content=md_table,
            metadata={"source": norm_filename, "page": 1, "type": "table"}
        ))
    except Exception as e:
        logger.error(f"Spreadsheet Error: {e}")
    return docs

def load_txt(path: str, filename: str) -> List[Document]:
    """
    Reads a plain-text file and returns a single text Document.
    Decodes with UTF-8 and ``errors='ignore'`` so Latin-1 or Windows-1252
    encoded files don't silently fail with UnicodeDecodeError.
    """
    docs = []
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            docs.append(Document(page_content=clean_text(f.read()), metadata={"source": normalize_source_key(filename), "page": 1, "type": "text"}))
    except Exception as e:
        logger.error(f"TXT Error: {e}")
    return docs

def load_docx(path: str, filename: str) -> List[Document]:
    """
    Extracts text from a Word document using XML-level ZIP parsing.
    Captures body paragraphs, table cells, text boxes, headers, footers,
    and SmartArt labels that ``python-docx``'s ``doc.paragraphs`` misses.
    Returns a single text Document (table structure in DOCX is not
    reconstructed as Markdown; paragraphs are emitted as plain text).
    """
    docs = []
    try:
        with open(path, 'rb') as f:
            file_bytes = f.read()
        full_text = clean_text(extract_docx_text(file_bytes))
        if full_text:
            docs.append(Document(
                page_content=full_text,
                metadata={"source": normalize_source_key(filename), "page": 1, "type": "text"}
            ))
    except Exception as e:
        logger.error(f"DOCX Error: {e}")
    return docs

def load_image(path: str, filename: str) -> List[Document]:
    """
    Extracts text from an image file (PNG/JPG/TIFF) via Tesseract OCR.
    Image is pre-processed with Otsu binarization before OCR to improve
    accuracy on low-contrast scans.  PSM 6 assumes a single uniform block
    of text.  Returns a text Document only if > 10 chars were extracted.
    """
    docs = []
    try:
        img = Image.open(path)
        clean_img = preprocess_image_for_ocr(img)
        ocr_text = pytesseract.image_to_string(clean_img, config="--psm 6 --oem 3")
        if len(ocr_text.strip()) > 10:
            docs.append(Document(page_content=clean_text(ocr_text), metadata={"source": normalize_source_key(filename), "page": 1, "type": "text"}))
    except Exception as e:
        logger.error(f"Image OCR Error: {e}")
    return docs

# --- INGESTION PIPELINE ---
def upload_in_batches(vectorstore, chunks, batch_size=50):
    """
    Uploads Document chunks to Pinecone in batches of ``batch_size`` to stay
    within API payload limits.  Raises on the first failed batch so that the
    caller (``ingest_all_files`` / ``ingest_uploaded_files``) can abort and
    preserve the source files for retry rather than silently partial-indexing.
    """
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i+batch_size]
        try:
            vectorstore.add_documents(batch)
            logger.info(f"Uploaded batch {i//batch_size + 1}")
        except Exception as e:
            logger.error(f"Batch {i//batch_size + 1} failed: {e}")
            raise 


def split_table_by_rows(doc: Document, max_rows: int = 20) -> List[Document]:
    """
    Splits a large Markdown table Document into sub-table chunks of max_rows each.
    Preserves the header row in every chunk so each is self-contained.
    Does NOT set chunk_index — the caller's source_counters loop handles that.
    Returns the original doc unchanged if it fits within max_rows or has no
    valid Markdown table structure.
    """
    lines = [l for l in doc.page_content.split('\n') if l.strip()]

    header_lines = []
    data_lines = []
    found_separator = False

    for line in lines:
        if not found_separator:
            header_lines.append(line)
            if line.strip().startswith('|') and '---' in line:
                found_separator = True
        else:
            data_lines.append(line)

    # Table fits within limit or no valid Markdown table structure found
    if len(data_lines) <= max_rows or not found_separator:
        return [doc]

    header = '\n'.join(header_lines)
    chunks = []
    for i in range(0, len(data_lines), max_rows):
        batch = data_lines[i:i + max_rows]
        chunk_content = header + '\n' + '\n'.join(batch)
        chunks.append(Document(
            page_content=chunk_content,
            metadata=dict(doc.metadata)  # Copy metadata; chunk_index set by source_counters
        ))
    return chunks


def ingest_all_files():
    """
    Scans all files under DOCS_FOLDER recursively, loads them into LangChain
    Document objects, chunks text docs, uploads everything to Pinecone, updates
    the Supabase manifest ledger, invalidates the semantic cache, and deletes
    the source files from disk to conserve storage.

    Returns:
        True  — all files successfully uploaded (or nothing new to ingest).
        False — the Pinecone upload step raised an exception; local files are
                preserved so the admin can retry.
    """
    if not os.path.exists(DOCS_FOLDER):
        os.makedirs(DOCS_FOLDER)
        return False

    all_docs = []
    logger.info(f"📂 Scanning {DOCS_FOLDER} and all subfolders...")

    # Phase 1: Recursive Scanning & Loading
    for root, dirs, files in os.walk(DOCS_FOLDER):
        rel_path = os.path.relpath(root, DOCS_FOLDER)
        raw_category = "general" if rel_path == "." else rel_path.split(os.sep)[0].lower()
    
        # Validate Category against constants
        if raw_category != "general" and raw_category not in VALID_CATEGORIES:
            logger.warning(f"⚠️ Unknown folder '{raw_category}' detected. Tagging as 'general'.")
            category = "general"
        else:
            category = raw_category

        # FIXED: This loop is now outside the 'else' so 'general' files are actually processed
        for filename in files:
            # NOTE: Duplicate detection uses bare filename as the source key.
            # Files in different folders with the same name are treated as the same
            # document.  A future improvement is to key by relative path + content hash.
            if is_already_ingested(filename):
                logger.info(f"Skipping {filename} — active in Pinecone.")
                continue
            
            file_path = os.path.join(root, filename)
            ext = filename.lower().split('.')[-1]
            
            file_docs = []
            if ext == 'pdf': file_docs = load_pdf(file_path, filename)
            elif ext in ['xlsx', 'xls', 'csv']: file_docs = load_spreadsheet(file_path, filename, ext == 'csv')
            elif ext == 'txt': file_docs = load_txt(file_path, filename)
            elif ext in ['doc', 'docx']: file_docs = load_docx(file_path, filename)
            elif ext in ['jpg', 'jpeg', 'png']: file_docs = load_image(file_path, filename)

            # Inject category and timestamp into metadata only at this stage.
            # The [CATEGORY | filename] context header is added to page_content
            # AFTER splitting (see Phase 2) so every split chunk carries it —
            # not just the first one.
            for d in file_docs:
                d.metadata["category"] = category
                d.metadata["uploaded_at"] = int(datetime.utcnow().timestamp())

            all_docs.extend(file_docs)

    if not all_docs: return True

    # Phase 2: Atomic Chunking
    table_docs = [d for d in all_docs if d.metadata.get("type") == "table"]
    text_docs  = [d for d in all_docs if d.metadata.get("type") == "text"]

    # Split large tables row-by-row so every chunk stays within the
    # all-MiniLM-L6-v2 256 word-piece embedding limit.
    split_table_docs = []
    for doc in table_docs:
        split_table_docs.extend(split_table_by_rows(doc, max_rows=20))

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    split_text_docs = text_splitter.split_documents(text_docs)
    final_chunks = split_table_docs + split_text_docs

    # ── Context Header Injection ───────────────────────────────────────────
    # CRITICAL: all-MiniLM-L6-v2 only embeds page_content — metadata is
    # invisible to it.  Without a context header, a chunk containing a
    # faculty list table (raw names, ranks, emails) won't semantically match
    # queries like "who are the ECE CpE faculty" because those words don't
    # appear in the table rows themselves.
    #
    # Injecting "[MEMOS] Faculty List ECE CpE" at the top of EVERY final chunk
    # (after splitting, so all chunks carry it — not just the first) lets the
    # embedding model link query terms to the correct document.
    #
    # Curriculum chunks are exempt: load_pdf already prepends
    # "Program: BACHELOR OF SCIENCE IN ..." for those files.
    for chunk in final_chunks:
        if not chunk.page_content.startswith("Program:"):
            src = chunk.metadata.get("source", "")
            category_label = chunk.metadata.get("category", "general")
            # Strip extension; replace underscores/hyphens with spaces so the
            # embedding tokeniser treats them as readable words.
            src_label = src.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")
            chunk.page_content = (
                f"[{category_label.upper()}] {src_label}\n\n{chunk.page_content}"
            )

    chunk_counts = {}
    source_counters = defaultdict(int)
    for chunk in final_chunks:
        src = chunk.metadata.get("source", "unknown")
        chunk.metadata["chunk_index"] = source_counters[src]
        source_counters[src] += 1
        chunk_counts[src] = chunk_counts.get(src, 0) + 1

    # Phase 3: Upload & Cleanup
    try:
        vectorstore = get_vectorstore()
        logger.info(f"🚀 Uploading {len(final_chunks)} chunks to Pinecone...")
        
        # 1. Batched Upload
        upload_in_batches(vectorstore, final_chunks)
        logger.info("✅ Ingestion Complete!")
        
        # 2. Update Manifest Ledger
        for filename, count in chunk_counts.items():
            update_manifest(filename, count)

        # 3. Automatic Cache Wipe (Ensures immediate AI updates)
        invalidate_cache()
            
        # 4. Delete source files whose keys appear in chunk_counts to free disk space.
        # chunk_counts is keyed by bare filename (the value passed as 'filename' to
        # each load_* function).  We match against the same bare filename here so
        # the lookup is correct.
        for root, dirs, files in os.walk(DOCS_FOLDER):
            for filename in files:
                if normalize_source_key(filename) in chunk_counts:
                    file_path = os.path.join(root, filename)
                    try:
                        os.remove(file_path)
                        logger.info(f"🗑️ Storage Optimized: Removed {filename}")
                    except Exception as e:
                        logger.error(f"⚠️ Could not delete {filename}: {e}")
                        
        return True
        
    except Exception as e:
        logger.error(f"❌ Upload Failed. Local files preserved for retry. Error: {e}")
        return False


# --- LEDGER ---
def get_uploaded_files() -> dict:
    """Reads manifest from Supabase instead of local JSON."""
    try:
        response = supabase.table("manifest").select("*").execute()
        return {
            row["filename"]: {
                "chunks": row["chunks"],
                "status": row["status"],
                "uploaded_at": str(row["uploaded_at"])
            }
            for row in response.data
        }
    except Exception as e:
        logger.error(f"Failed to read manifest from Supabase: {e}")
        return {}

def update_manifest(filename: str, chunk_count: int):
    """Upserts a file record into Supabase manifest."""
    norm_filename = normalize_source_key(filename)
    try:
        supabase.table("manifest").upsert({
            "filename": norm_filename,
            "chunks": chunk_count,
            "status": "Active",
            "uploaded_at": datetime.utcnow().isoformat()
        }).execute()
    except Exception as e:
        logger.error(f"Failed to update manifest: {e}")

def remove_from_manifest(filename: str):
    """Deletes a file record from Supabase manifest."""
    norm_filename = normalize_source_key(filename)
    try:
        supabase.table("manifest").delete().eq("filename", norm_filename).execute()
    except Exception as e:
        logger.error(f"Failed to remove from manifest: {e}")

def delete_document(filename: str) -> Tuple[bool, str]:
    """
    Deletes all Pinecone vectors whose ``source`` metadata matches ``filename``
    and removes the corresponding entry from the Supabase manifest.

    Args:
        filename: The bare filename or normalized relative path as stored in
                  the manifest (e.g. ``"Faculty_List.pdf"`` or
                  ``"curriculum/BS_CpE.pdf"``).

    Returns:
        (True, success_message) or (False, error_message).
    """
    norm_filename = normalize_source_key(filename)
    try:
        pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
        index = pc.Index(PINECONE_INDEX_NAME)
        index.delete(filter={"source": norm_filename})
        remove_from_manifest(norm_filename)
        return True, f"Successfully purged {norm_filename} chunks."
    except Exception as e:
        return False, f"Failed to delete {norm_filename}: {str(e)}"

def purge_all_vectors() -> Tuple[bool, str]:
    """Nuclear option: wipe ALL vectors from Pinecone and clear the manifest."""
    try:
        pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
        index = pc.Index(PINECONE_INDEX_NAME)
        index.delete(delete_all=True)
        # Clear entire manifest table
        try:
            supabase.table("manifest").delete().neq("filename", "").execute()
        except Exception:
            pass  # Manifest may already be empty
        logger.info("🗑️ Purged ALL vectors from Pinecone and cleared manifest.")
        return True, "All vectors deleted from Pinecone. Index is now empty."
    except Exception as e:
        return False, f"Purge failed: {str(e)}"
    
def verify_sync() -> dict:
    """
    Compares the Supabase manifest against source keys visible in Pinecone.

    KNOWN LIMITATION: Pinecone does not expose a full enumeration API on the
    serverless/free tier.  This function approximates the index contents by
    issuing an ANN query against a zero vector, which returns the 10,000 nearest
    neighbors to the origin — NOT a complete scan.  Treat results as indicative,
    not authoritative.  Do not make irreversible decisions (e.g., manifest
    cleanup) based solely on this output.
    """
    manifest = get_uploaded_files()
    manifest_sources = set(manifest.keys())
    
    try:
        pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
        index = pc.Index(PINECONE_INDEX_NAME)
        
        # We fetch unique 'source' values from the index
        # Note: Depending on Pinecone version, you may need to use list_ids or a dummy query
        results = index.query(vector=[0]*384, top_k=10000, include_metadata=True)
        pinecone_sources = {d['metadata']['source'] for d in results['matches']}
        
        return {
            "in_both": list(manifest_sources & pinecone_sources),
            "manifest_only": list(manifest_sources - pinecone_sources), # Ghost entries
            "pinecone_only": list(pinecone_sources - manifest_sources)  # Untracked vectors
        }
    except Exception as e:
        logger.error(f"Sync check failed: {e}")
        return {}
    
def check_pinecone_health() -> bool:
    """Checks Pinecone connectivity and index availability via describe_index_stats."""
    try:
        from pinecone import Pinecone
        pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
        index = pc.Index(PINECONE_INDEX_NAME)
        stats = index.describe_index_stats()
        logger.info(f"✅ Pinecone healthy: {stats.total_vector_count} vectors indexed")
        return True
    except Exception as e:
        logger.error(f"🚨 Pinecone Health Check Failed: {e}")
        return False